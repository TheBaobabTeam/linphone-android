from docx import Document
from docx.shared import Inches
import xml.etree.ElementTree as ET


def parseXML():
	#tree = ET.parse('report.xml')
	#root = tree.getroot()

	#get the document
	document = Document()

	document.add_heading('Unit Testing Results Document', 0)

	details_table = document.add_table(rows=5, cols=4)

	p_name = details_table.cell(0,0).text = 'Project Name: '
	version = details_table.cell(0,2).text = 'Version:' 
	author = details_table.cell(1,0).text = 'Written by: '
	desc = details_table.cell(1,2).text = 'Description '

	document.add_heading('Table showing results', level=2)

	test_table = document.add_table(rows=5, cols=7)

	# populate header row --------
	heading_cells = test_table.rows[0].cells
	heading_cells[0].text = '#'
	heading_cells[1].text = 'Date Added'
	heading_cells[2].text = 'Use case'
	heading_cells[3].text = 'Expected Result'
	heading_cells[4].text = 'Passed?'
	heading_cells[5].text = 'Failed?'
	heading_cells[6].text = 'Remarks'

	input_cell = test_table.rows[1].cells

	input_cell[0].text = '1'
	input_cell[1].text = '7 September 2015'
	input_cell[2].text = 'Create Group Chat'
	
	document.save('Acceptance_Test_Results.docx')




parseXML()


